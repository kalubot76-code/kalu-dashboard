"""
Kalu Advanced Document Generator
=================================

Gera documentos em múltiplos formatos: HTML, PDF, Word, Excel
"""

import json
import os
from datetime import datetime
from typing import Dict, Any, Optional
import base64

def generate_word_document(data: Dict[str, Any], task_title: str, output_path: str) -> str:
    """
    Gera documento Word (.docx)
    
    Args:
        data: Dados do relatório
        task_title: Título da tarefa
        output_path: Caminho completo para salvar
        
    Returns:
        str: Caminho do ficheiro gerado
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Título principal
        title = doc.add_heading(task_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(31, 119, 180)  # Azul
        
        # Data e info
        info = doc.add_paragraph()
        info.add_run(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n").bold = True
        info.add_run(f"Gerado por: Kalu AI Assistant ⚡")
        
        doc.add_paragraph()  # Espaço
        
        # Processar dados
        if isinstance(data, str):
            try:
                data = json.loads(data)
            except:
                doc.add_paragraph(data)
                doc.save(output_path)
                return output_path
        
        # Extrair estrutura (ex: análise_concorrentes, plano_expansao)
        main_key = list(data.keys())[0] if data else None
        content = data.get(main_key, {}) if main_key else data
        
        # Iterar pelos campos e formatar
        for key, value in content.items():
            if key in ['data', 'empresa', 'mercado', 'segmento']:
                # Informações gerais
                p = doc.add_paragraph()
                p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                p.add_run(str(value))
            
            elif isinstance(value, list):
                # Listas
                heading = doc.add_heading(key.replace('_', ' ').title(), 2)
                for item in value:
                    if isinstance(item, dict):
                        # Item complexo
                        for subkey, subvalue in item.items():
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(f"{subkey.replace('_', ' ').title()}: ").bold = True
                            p.add_run(str(subvalue))
                    else:
                        doc.add_paragraph(str(item), style='List Bullet')
            
            elif isinstance(value, dict):
                # Subsecções
                heading = doc.add_heading(key.replace('_', ' ').title(), 2)
                for subkey, subvalue in value.items():
                    if isinstance(subvalue, (list, dict)):
                        # Recursivo
                        subheading = doc.add_heading(subkey.replace('_', ' ').title(), 3)
                        p = doc.add_paragraph(str(subvalue))
                    else:
                        p = doc.add_paragraph()
                        p.add_run(f"{subkey.replace('_', ' ').title()}: ").bold = True
                        p.add_run(str(subvalue))
            
            elif key in ['conclusao', 'visao']:
                # Conclusão destacada
                doc.add_heading('Conclusão', 2)
                doc.add_paragraph(str(value))
        
        # Rodapé
        doc.add_page_break()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("Gerado automaticamente por Kalu AI Assistant")
        footer_run.italic = True
        footer_run.font.size = Pt(10)
        footer_run.font.color.rgb = RGBColor(127, 140, 141)
        
        doc.save(output_path)
        return output_path
        
    except ImportError:
        raise Exception("Biblioteca python-docx não instalada. Execute: pip install python-docx")


def generate_excel_document(data: Dict[str, Any], task_title: str, output_path: str) -> str:
    """
    Gera documento Excel (.xlsx) com dados tabulares
    
    Args:
        data: Dados do relatório
        task_title: Título da tarefa
        output_path: Caminho completo para salvar
        
    Returns:
        str: Caminho do ficheiro gerado
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill
        from openpyxl.utils import get_column_letter
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Relatório"
        
        # Título
        ws['A1'] = task_title
        ws['A1'].font = Font(size=16, bold=True, color="1F77B4")
        ws['A1'].alignment = Alignment(horizontal='center')
        ws.merge_cells('A1:E1')
        
        # Data
        ws['A2'] = f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        ws['A2'].font = Font(italic=True)
        
        row = 4  # Começar dados na linha 4
        
        # Processar dados
        if isinstance(data, str):
            try:
                data = json.loads(data)
            except:
                ws.cell(row=row, column=1, value=data)
                wb.save(output_path)
                return output_path
        
        # Extrair estrutura
        main_key = list(data.keys())[0] if data else None
        content = data.get(main_key, {}) if main_key else data
        
        # Criar tabelas por secção
        for key, value in content.items():
            # Cabeçalho de secção
            ws.cell(row=row, column=1, value=key.replace('_', ' ').title())
            ws.cell(row=row, column=1).font = Font(size=14, bold=True, color="2C3E50")
            ws.cell(row=row, column=1).fill = PatternFill(start_color="E8F4F8", end_color="E8F4F8", fill_type="solid")
            row += 1
            
            if isinstance(value, list):
                # Tabela de lista
                for i, item in enumerate(value, 1):
                    if isinstance(item, dict):
                        if i == 1:
                            # Cabeçalhos
                            for col, header in enumerate(item.keys(), 1):
                                cell = ws.cell(row=row, column=col, value=header.replace('_', ' ').title())
                                cell.font = Font(bold=True)
                                cell.fill = PatternFill(start_color="D6EAF8", end_color="D6EAF8", fill_type="solid")
                            row += 1
                        
                        # Dados
                        for col, val in enumerate(item.values(), 1):
                            ws.cell(row=row, column=col, value=str(val))
                        row += 1
                    else:
                        ws.cell(row=row, column=1, value=str(item))
                        row += 1
            
            elif isinstance(value, dict):
                # Tabela chave-valor
                for subkey, subvalue in value.items():
                    ws.cell(row=row, column=1, value=subkey.replace('_', ' ').title())
                    ws.cell(row=row, column=1).font = Font(bold=True)
                    ws.cell(row=row, column=2, value=str(subvalue))
                    row += 1
            
            else:
                # Valor simples
                ws.cell(row=row, column=1, value=str(value))
                row += 1
            
            row += 1  # Espaço entre secções
        
        # Ajustar largura das colunas
        for column in range(1, 6):
            ws.column_dimensions[get_column_letter(column)].width = 25
        
        # Rodapé
        row += 2
        ws.cell(row=row, column=1, value="Gerado automaticamente por Kalu AI Assistant")
        ws.cell(row=row, column=1).font = Font(italic=True, size=9, color="7F8C8D")
        
        wb.save(output_path)
        return output_path
        
    except ImportError:
        raise Exception("Biblioteca openpyxl não instalada. Execute: pip install openpyxl")


def generate_pdf_from_html(html_content: str, output_path: str) -> str:
    """
    Converte HTML para PDF
    
    Args:
        html_content: Conteúdo HTML completo
        output_path: Caminho para salvar PDF
        
    Returns:
        str: Caminho do ficheiro gerado
    """
    try:
        # Opção 1: wkhtmltopdf (requer instalação sistema)
        import subprocess
        
        # Salvar HTML temporário
        temp_html = output_path.replace('.pdf', '_temp.html')
        with open(temp_html, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # Converter
        result = subprocess.run(
            ['wkhtmltopdf', '--enable-local-file-access', temp_html, output_path],
            capture_output=True,
            text=True
        )
        
        # Limpar temp
        if os.path.exists(temp_html):
            os.remove(temp_html)
        
        if result.returncode == 0:
            return output_path
        else:
            raise Exception(f"wkhtmltopdf error: {result.stderr}")
            
    except FileNotFoundError:
        # wkhtmltopdf não instalado, tentar alternativa
        try:
            # Opção 2: weasyprint (Python puro)
            from weasyprint import HTML
            HTML(string=html_content).write_pdf(output_path)
            return output_path
        except ImportError:
            raise Exception("wkhtmltopdf ou weasyprint não instalados. Instale um deles.")


def generate_document_with_downloads(
    data: Dict[str, Any],
    task_title: str,
    task_id: int,
    output_dir: str = "/tmp/kalu_reports"
) -> Dict[str, str]:
    """
    Gera documento em múltiplos formatos e retorna caminhos
    
    Args:
        data: Dados do relatório
        task_title: Título
        task_id: ID da tarefa
        output_dir: Directório de output
        
    Returns:
        dict: {"html": path, "docx": path, "xlsx": path, "pdf": path}
    """
    os.makedirs(output_dir, exist_ok=True)
    
    base_filename = f"task_{task_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    results = {}
    
    # HTML (sempre gera)
    from kalu_document_generator import generate_report
    html_content = generate_report(data, task_title, "html", content_only=True)
    results['html'] = html_content
    
    # Word
    try:
        docx_path = os.path.join(output_dir, f"{base_filename}.docx")
        generate_word_document(data, task_title, docx_path)
        results['docx'] = docx_path
    except Exception as e:
        print(f"⚠️ Erro ao gerar Word: {e}")
        results['docx'] = None
    
    # Excel
    try:
        xlsx_path = os.path.join(output_dir, f"{base_filename}.xlsx")
        generate_excel_document(data, task_title, xlsx_path)
        results['xlsx'] = xlsx_path
    except Exception as e:
        print(f"⚠️ Erro ao gerar Excel: {e}")
        results['xlsx'] = None
    
    # PDF (precisa HTML completo)
    try:
        from kalu_document_generator import generate_html_report, generate_markdown_report
        markdown = generate_markdown_report(data, task_title)
        html_full = generate_html_report(markdown)
        
        pdf_path = os.path.join(output_dir, f"{base_filename}.pdf")
        generate_pdf_from_html(html_full, pdf_path)
        results['pdf'] = pdf_path
    except Exception as e:
        print(f"⚠️ Erro ao gerar PDF: {e}")
        results['pdf'] = None
    
    return results


if __name__ == "__main__":
    # Teste
    sample_data = {
        "teste": {
            "titulo": "Documento de Teste",
            "items": ["Item 1", "Item 2", "Item 3"]
        }
    }
    
    try:
        results = generate_document_with_downloads(
            data=sample_data,
            task_title="Teste Multi-Formato",
            task_id=999
        )
        print("Documentos gerados:")
        for fmt, path in results.items():
            print(f"  {fmt}: {path}")
    except Exception as e:
        print(f"Erro: {e}")
